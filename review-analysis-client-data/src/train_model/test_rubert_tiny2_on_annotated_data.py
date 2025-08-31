import os
import sys
import logging
from datetime import datetime
import random
import numpy as np
import pandas as pd
import torch
import torch.nn.functional as F
from torch.utils.data import DataLoader

# Добавляем корень проекта в путь
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
sys.path.append(PROJECT_ROOT)

# Импорты HuggingFace
from datasets import Dataset
from transformers import BertTokenizer, BertForSequenceClassification

# Для отчётов
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Метрики
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report, roc_auc_score
from sklearn.metrics import confusion_matrix, roc_curve, auc

# Визуализация
import matplotlib.pyplot as plt
import seaborn as sns

# Лемматизация
from pymorphy3 import MorphAnalyzer
import re

# Логгирование
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("test_model")
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(logging.Formatter("[%(asctime)s] %(levelname)s: %(message)s"))
logger.addHandler(console_handler)

# Пути
LOGS_DIR = os.path.join(PROJECT_ROOT, "logs")
MODELS_DIR = os.path.join(PROJECT_ROOT, "models")
RESULTS_DIR = os.path.join(PROJECT_ROOT, "results")
os.makedirs(LOGS_DIR, exist_ok=True)
os.makedirs(MODELS_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)

TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M")
RESULT_CSV = os.path.join(RESULTS_DIR, f"predictions_test_class_7_and_other_{TIMESTAMP}.csv")
REPORT_DOCX = os.path.join(RESULTS_DIR, f"report_test_class_7_and_other_{TIMESTAMP}.docx")

# Гиперпараметры
MODEL_NAME = "cointegrated/rubert-tiny2"
MAX_LENGTH = 256
BATCH_SIZE = 16
SEED = 42

# Фиксируем seed
random.seed(SEED)
np.random.seed(SEED)
torch.manual_seed(SEED)

# Инициализируем лемматизатор
morph = MorphAnalyzer()

def preprocess_text(text):
    """Приводит текст к нижнему регистру, удаляет пунктуацию, лемматизирует"""
    text = str(text).lower()
    text = re.sub(r'[^\w\s]', '', text)  # Удаляем знаки препинания
    text = re.sub(r'\s+', ' ', text).strip()  # Удаляем лишние пробелы
    tokens = text.split()
    lemmatized = [morph.parse(token)[0].normal_form for token in tokens if token.isalpha()]
    return ' '.join(lemmatized)

def load_data():
    logger.info("📥 Загрузка данных...")
    DATA_PATH = os.path.join(PROJECT_ROOT, "data", "ALL_Reviews_P23.csv")
    df = pd.read_csv(DATA_PATH, sep=',', on_bad_lines='skip', encoding='utf-8', engine='python')

    if 'Класс' in df.columns and 'Текст отзыва' in df.columns:
        df = df.rename(columns={"Класс": "label", "Текст отзыва": "text"})
    else:
        logger.error("❌ Колонка 'Класс' или 'Текст отзыва' отсутствует.")
        print(df.head(3).to_string())
        raise KeyError("Формат файла некорректен или данные повреждены.")

    # Применяем предобработку и лемматизацию
    df['text'] = df['text'].astype(str).apply(preprocess_text)
    df['label'] = df['label'].apply(lambda x: 1 if x == 7 else 0)

    return df

def prepare_dataset(df):
    logger.info("🛠 Подготовка датасета...")

    # Сбалансированная выборка: по 2300 для каждого класса
    class_7 = df[df['label'] == 1].sample(n=2300, random_state=SEED)
    class_0 = df[df['label'] == 0].sample(n=2300, random_state=SEED)

    # Объединяем и перемешиваем
    combined = pd.concat([class_7, class_0]).sample(frac=1, random_state=SEED).reset_index(drop=True)

    logger.info(f"📊 Выборка: {len(combined)} отзывов, распределение классов:")
    print(combined['label'].value_counts())

    test_dataset = Dataset.from_pandas(combined[['text', 'label']])
    return test_dataset

def tokenize_dataset(dataset, tokenizer):
    def tokenize_function(examples):
        texts = [str(text) for text in examples["text"]]
        return tokenizer(
            texts,
            padding="max_length",
            truncation=True,
            max_length=MAX_LENGTH,
            return_tensors="pt"
        )

    return dataset.map(tokenize_function, batched=True, remove_columns=["text"])

def evaluate_model(model, tokenizer, test_dataset):
    logger.info("📊 Оценка модели на тестовой выборке...")

    tokenized_test = tokenize_dataset(test_dataset, tokenizer).with_format("torch")
    test_loader = DataLoader(tokenized_test, batch_size=BATCH_SIZE)

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)
    model.eval()

    true_labels = []
    predictions = []
    probs_list = []

    with torch.no_grad():
        for batch in test_loader:
            inputs = {k: v.to(device) for k, v in batch.items() if k in ["input_ids", "attention_mask"]}
            labels = batch["label"].to(device)

            outputs = model(**inputs)
            logits = outputs.logits
            probs = F.softmax(logits, dim=-1)
            preds = torch.argmax(logits, dim=-1)

            true_labels.extend(labels.cpu().numpy())
            predictions.extend(preds.cpu().numpy())
            probs_list.extend(probs.cpu().numpy())

    probs_array = np.array(probs_list)

    # Сохраняем результаты
    df_results = pd.DataFrame({
        "text": test_dataset["text"],
        "true_label": true_labels,
        "predicted_label": predictions,
        "prob_class_7": probs_array[:, 1],
        "prob_class_other": probs_array[:, 0]
    })

    df_results.to_csv(RESULT_CSV, index=False)
    logger.info(f"📁 Результаты сохранены в {RESULT_CSV}")

    # Метрики
    acc = accuracy_score(true_labels, predictions)
    f1 = f1_score(true_labels, predictions)
    prec = precision_score(true_labels, predictions)
    rec = recall_score(true_labels, predictions)
    roc_auc = roc_auc_score(true_labels, probs_array[:, 1])

    print(f"🎯 Метрики:")
    print(f"Accuracy: {acc:.4f}, F1: {f1:.4f}, Precision: {prec:.4f}, Recall: {rec:.4f}, ROC AUC: {roc_auc:.4f}")
    print(classification_report(true_labels, predictions, target_names=["Класс 0", "Класс 7"]))

    return df_results, true_labels, predictions, probs_array

def plot_results(true_labels, probs, save_dir, timestamp):
    logger.info("📉 Построение графиков...")

    # Матрица ошибок
    cm = confusion_matrix(true_labels, np.argmax(probs, axis=1))
    plt.figure(figsize=(6, 6))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", xticklabels=["Класс 0", "Класс 7"],
                yticklabels=["Класс 0", "Класс 7"])
    plt.title("Confusion Matrix")
    plt.tight_layout()
    cm_path = os.path.join(save_dir, f"confusion_matrix_class_7_and_other_{timestamp}.png")
    plt.savefig(cm_path)
    plt.close()

    # Гистограмма вероятностей
    plt.figure(figsize=(10, 6))
    plt.hist(probs[np.array(true_labels) == 1, 1], bins=50, alpha=0.7, label="Класс 7", color="green")
    plt.hist(probs[np.array(true_labels) == 0, 1], bins=50, alpha=0.7, label="Класс 0", color="red")
    plt.title("Распределение вероятностей для классов")
    plt.xlabel("Вероятность быть классом 7")
    plt.ylabel("Частота")
    plt.legend()
    hist_path = os.path.join(save_dir, f"probability_histogram_class_7_and_other_{timestamp}.png")
    plt.savefig(hist_path)
    plt.close()

    # ROC-кривая
    fpr, tpr, _ = roc_curve(true_labels, probs[:, 1])
    roc_auc = auc(fpr, tpr)
    plt.figure(figsize=(8, 6))
    plt.plot(fpr, tpr, color='darkorange', lw=2, label=f'ROC curve (AUC = {roc_auc:.2f})')
    plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
    plt.xlim([0.0, 1.0])
    plt.ylim([0.0, 1.05])
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curve')
    plt.legend(loc="lower right")
    roc_path = os.path.join(save_dir, f"roc_curve_class_7_and_other_{timestamp}.png")
    plt.savefig(roc_path)
    plt.close()

    return cm_path, hist_path, roc_path

def save_report_to_docx(df_results, true_labels, predictions, probs, cm_path, hist_path, roc_path, save_dir, timestamp):
    logger.info("📄 Создание отчёта в Word (.docx)...")

    doc = Document()
    doc.add_heading("Отчёт по тестированию модели: Класс 7 vs Класс 0", level=1)
    doc.add_paragraph(f"Дата и время: {timestamp}")

    # Метрики
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Метрика"
    hdr_cells[1].text = "Значение"

    metrics = {
        "accuracy": accuracy_score(true_labels, predictions),
        "f1": f1_score(true_labels, predictions),
        "precision": precision_score(true_labels, predictions),
        "recall": recall_score(true_labels, predictions),
        "roc_auc": roc_auc_score(true_labels, probs[:, 1]),
    }

    for key, value in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = f"{value:.4f}"

    # Классификационный отчёт
    doc.add_heading("Классификационный отчёт", level=2)
    report = classification_report(true_labels, predictions, target_names=["Класс 0", "Класс 7"], digits=4)
    doc.add_paragraph(report)

    # Примеры
    doc.add_heading("Примеры предсказаний", level=2)
    sample = df_results.head(10)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Текст"
    hdr_cells[1].text = "Метка"
    hdr_cells[2].text = "Предсказание"
    hdr_cells[3].text = "Вероятность (Класс 7)"
    hdr_cells[4].text = "Вероятность (Класс 0)"
    for _, row in sample.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['text'][:100] + "..." if len(row['text']) > 100 else row['text']
        row_cells[1].text = str(row['true_label'])
        row_cells[2].text = str(row['predicted_label'])
        row_cells[3].text = f"{row['prob_class_7']:.4f}"
        row_cells[4].text = f"{row['prob_class_other']:.4f}"

    # Графики
    doc.add_heading("Графики", level=2)
    doc.add_picture(cm_path, width=Mm(150))
    doc.add_paragraph("Матрица ошибок", style='Caption')
    doc.add_picture(hist_path, width=Mm(150))
    doc.add_paragraph("Гистограмма вероятностей", style='Caption')
    doc.add_picture(roc_path, width=Mm(150))
    doc.add_paragraph("ROC-кривая", style='Caption')

    # Сохраняем отчёт
    report_path = os.path.join(save_dir, f"report_test_class_7_and_other_{timestamp}.docx")
    doc.save(report_path)
    logger.info(f"✅ Отчёт сохранён в {report_path}")
    return report_path

if __name__ == "__main__":
    logger.info("🚀 Начало тестирования модели на новых данных")

    # Загружаем данные и лемматизируем
    df = load_data()

    # Подготавливаем датасет
    test_dataset = prepare_dataset(df)

    # Загружаем модель
    #MODEL_SAVE_DIR = r"./models/model_class_7_and_0_20250721_1127"
    MODEL_SAVE_DIR = r"./models/model_class_7_and_0_20250728_0702"
    logger.info(f"📥 Загрузка модели из {MODEL_SAVE_DIR}")

    try:
        tokenizer = BertTokenizer.from_pretrained(MODEL_NAME)
        model = BertForSequenceClassification.from_pretrained(MODEL_SAVE_DIR, num_labels=2)
    except Exception as e:
        logger.error(f"❌ Не удалось загрузить модель: {e}")
        raise

    # Тестируем модель
    df_results, true_labels, predictions, probs = evaluate_model(model, tokenizer, test_dataset)

    # Строим графики
    cm_path, hist_path, roc_path = plot_results(true_labels, probs, RESULTS_DIR, TIMESTAMP)

    # Сохраняем отчёт в Word
    save_report_to_docx(df_results, true_labels, predictions, probs, cm_path, hist_path, roc_path, RESULTS_DIR, TIMESTAMP)

    logger.info("✅ Тестирование завершено успешно!")