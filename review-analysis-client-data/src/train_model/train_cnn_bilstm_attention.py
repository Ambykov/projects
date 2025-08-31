import os
import re
import random
from collections import Counter
import logging
from datetime import datetime

import numpy as np
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.metrics import (accuracy_score, f1_score, precision_score, recall_score,
                             roc_auc_score, confusion_matrix, classification_report, roc_curve, auc)

import torch
import torch.nn as nn
import torch.optim as optim
from torch.utils.data import Dataset, DataLoader

import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Mm

from gensim.models import KeyedVectors
import pymorphy3
import pickle


# --- НАСТРОЙКИ ПУТЕЙ ---
PROJECT_ROOT = r"C:\Users\Андрей\Documents\Курсы\AI_ML_разработчик\Стажировка_2\Ноутбуки\Пайплайн"
DATA_PATH = os.path.join(PROJECT_ROOT, "data", "full_reviews_mixed.csv")
EMBEDDING_PATH = r"C:\Users\Андрей\Documents\Курсы\AI_ML_разработчик\Стажировка_2\ruwikiruscorpora\model.txt"

RESULTS_DIR = os.path.join(PROJECT_ROOT, "results", "cnn_bilstm_attention_class_7")
os.makedirs(RESULTS_DIR, exist_ok=True)

TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M")
REPORT_PATH = os.path.join(RESULTS_DIR, f"report_{TIMESTAMP}.docx")
RESULT_CSV = os.path.join(RESULTS_DIR, f"predictions_{TIMESTAMP}.csv")
MODEL_PATH = os.path.join(RESULTS_DIR, f"model_{TIMESTAMP}.pt")
MISSED_WORDS_PATH = os.path.join(RESULTS_DIR, f"missed_words_{TIMESTAMP}.txt")
WORD2IDX_PATH = os.path.join(RESULTS_DIR, "word2idx.pkl")
EMBED_MATRIX_PATH = os.path.join(RESULTS_DIR, "embed_matrix.npy")

# --- ГИПЕРПАРАМЕТРЫ ---
SEED = 42
MAX_LENGTH = 258
VOCAB_SIZE = 25000
EMBEDDING_DIM = 300
HIDDEN_SIZE = 128
NUM_CLASSES = 2
CNN_KERNELS = (3, 4, 5)
CNN_FILTERS = 64
BATCH_SIZE = 32
EPOCHS = 10
LEARNING_RATE = 2e-4
DROPOUT = 0.4
TOP_N_MISSED = 500

random.seed(SEED)
np.random.seed(SEED)
torch.manual_seed(SEED)
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("CNN_BiLSTM_Attention")


# --- ЛЕММАТИЗАЦИЯ ---
morph = pymorphy3.MorphAnalyzer()
pos_map = {
    'NOUN': 'NOUN',  'VERB': 'VERB', 'INFN': 'VERB',
    'ADJF': 'ADJ',   'ADJS': 'ADJ',  'COMP': 'ADJ',
    'PRTF': 'ADJ',   'PRTS': 'ADJ',  'GRND': 'VERB',
    'NUMR': 'NUM',   'ADVB': 'ADV',  'NPRO': 'PRON',
    'PRED': 'ADV',   'PREP': 'ADP',  'CONJ': 'CONJ',
    'PRCL': 'PART',  'INTJ': 'INTJ'
}

def get_pos_tag(word):
    parse = morph.parse(word)[0]
    pos = parse.tag.POS
    return pos_map.get(pos, 'X')

def process_token(token):
    parse = morph.parse(token)[0]
    lemma = parse.normal_form
    pos = get_pos_tag(token)
    return f"{lemma}_{pos}"

def preproc_text(text):
    tokens = re.findall(r'\w+', text.lower())
    return ' '.join(process_token(token) for token in tokens)


# --- ЗАГРУЗКА ПРЕДОБУЧЕННЫХ ЭМБЕДДИНГОВ И ФОРМИРОВАНИЕ СЛОВАРЯ ---
def build_vocab_and_embeddings(texts, vocab_size, embedding_path, embedding_dim):
    logger.info("Загрузка предобучённых эмбеддингов (это может занять время)...")
    w2v = KeyedVectors.load_word2vec_format(embedding_path, binary=False)
    counter = Counter()
    for text in texts:
        tokens = text.split()
        counter.update(tokens)
    most_common = counter.most_common(vocab_size - 2)
    word2idx = {"<PAD>": 0, "<UNK>": 1}
    word2idx.update({w: i + 2 for i, (w, _) in enumerate(most_common)})

    embed_matrix = np.zeros((len(word2idx), embedding_dim), dtype=np.float32)
    embed_matrix[1] = np.mean(w2v.vectors, axis=0)
    missed_words = []
    for word, idx in word2idx.items():
        if idx < 2:
            continue
        if word in w2v:
            embed_matrix[idx] = w2v[word]
        else:
            embed_matrix[idx] = embed_matrix[1]
            missed_words.append(word)

    logger.info(f"Embedding - неизвестных векторов: {len(missed_words)} из {len(word2idx)}")

    missed_counter = {w: counter[w] for w in missed_words}
    top_missed = sorted(missed_counter.items(), key=lambda x: x[1], reverse=True)
    print(f"\nТоп-{TOP_N_MISSED} не попавших в embedding слов (токен: частота в train):")
    for word, freq in top_missed[:TOP_N_MISSED]:
        print(f"{word}: {freq}")
    with open(MISSED_WORDS_PATH, "w", encoding="utf-8") as f:
        f.write("Токен\tЧастота появления в train\n")
        for word, freq in top_missed:
            f.write(f"{word}\t{freq}\n")
    logger.info(f"Полный список не вошедших в embedding слов сохранён в {MISSED_WORDS_PATH}")

    # Сохраняем словарь и embedding matrix
    with open(WORD2IDX_PATH, "wb") as f:
        pickle.dump(word2idx, f)
    np.save(EMBED_MATRIX_PATH, embed_matrix)
    logger.info(f"Словарь сохранён: {WORD2IDX_PATH}")
    logger.info(f"Матрица эмбеддингов сохранена: {EMBED_MATRIX_PATH}")

    return word2idx, embed_matrix


# --- ТРАНСФОРМАЦИЯ ТЕКСТА В ПОСЛЕДОВАТЕЛЬНОСТЬ ИНДЕКСОВ ---
def text_to_sequence(text, word2idx, max_length=MAX_LENGTH):
    tokens = text.split()
    indices = [word2idx.get(t, 1) for t in tokens]  # 1 - <UNK>
    if len(indices) < max_length:
        indices += [0] * (max_length - len(indices))
    else:
        indices = indices[:max_length]
    return indices


class ReviewDataset(Dataset):
    def __init__(self, texts, labels, word2idx, max_length=MAX_LENGTH):
        self.sequences = [text_to_sequence(txt, word2idx, max_length) for txt in texts]
        self.labels = labels

    def __len__(self):
        return len(self.sequences)

    def __getitem__(self, idx):
        return torch.tensor(self.sequences[idx], dtype=torch.long), torch.tensor(self.labels[idx], dtype=torch.long)


# --- CNN-BiLSTM-ATTENTION МОДЕЛЬ ---
class CNN_BiLSTM_Attention(nn.Module):
    def __init__(self, vocab_size, embedding_dim, hidden_size, num_classes,
                 kernel_sizes=(3,4,5), num_filters=64, dropout_rate=0.4, embed_matrix=None):
        super().__init__()
        if embed_matrix is not None:
            self.embedding = nn.Embedding.from_pretrained(torch.FloatTensor(embed_matrix), freeze=False, padding_idx=0)
        else:
            self.embedding = nn.Embedding(vocab_size, embedding_dim, padding_idx=0)
        self.convs = nn.ModuleList([
            nn.Conv1d(embedding_dim, num_filters, k) for k in kernel_sizes
        ])
        self.lstm = nn.LSTM(num_filters * len(kernel_sizes), hidden_size, batch_first=True, bidirectional=True)
        self.dropout = nn.Dropout(dropout_rate)
        self.attn_fc = nn.Linear(hidden_size * 2, 1)
        self.fc = nn.Sequential(
            nn.Linear(hidden_size * 2, hidden_size),
            nn.ReLU(),
            nn.Dropout(dropout_rate),
            nn.Linear(hidden_size, num_classes)
        )

    def forward(self, x):
        emb = self.embedding(x)  # [batch, seq_len, emb_dim]
        emb = emb.transpose(1, 2)  # [batch, emb_dim, seq_len]
        convs = [torch.relu(conv(emb)) for conv in self.convs]
        min_len = min(conv.shape[2] for conv in convs)
        convs = [conv[:, :, :min_len] for conv in convs]
        cnn_out = torch.cat(convs, dim=1)  # [batch, num_filters*len_kernels, seq_len']
        cnn_out = cnn_out.transpose(1, 2)  # [batch, seq_len', num_filters*len_kernels]
        lstm_out, _ = self.lstm(cnn_out)  # [batch, seq_len', hidden_size*2]
        attn_weights = torch.softmax(self.attn_fc(lstm_out), dim=1)  # [batch, seq_len', 1]
        attended = torch.sum(attn_weights * lstm_out, dim=1)        # [batch, hidden_size*2]
        out = self.dropout(attended)
        return self.fc(out)


# --- ТРЕНИРОВКА ---
def train_model(model, train_loader, criterion, optimizer, device):
    model.train()
    total_loss = 0
    for x_batch, y_batch in train_loader:
        x_batch = x_batch.to(device)
        y_batch = y_batch.to(device)
        optimizer.zero_grad()
        logits = model(x_batch)
        loss = criterion(logits, y_batch)
        loss.backward()
        optimizer.step()
        total_loss += loss.item()
    return total_loss / len(train_loader)


# --- ОЦЕНКА ---
def evaluate_model(model, loader, device):
    model.eval()
    preds, probs, trues = [], [], []
    with torch.no_grad():
        for x, y in loader:
            x = x.to(device)
            y = y.to(device)
            logits = model(x)
            prob = torch.softmax(logits, -1).cpu().numpy()
            pred = prob.argmax(-1)
            preds.extend(pred)
            probs.extend(prob)
            trues.extend(y.cpu().numpy())
    return np.array(trues), np.array(preds), np.array(probs)


# --- ГРАФИКИ И ОТЧЁТ ---
def plot_and_report(trues, preds, probs, texts, save_dir, timestamp):
    acc = accuracy_score(trues, preds)
    f1 = f1_score(trues, preds)
    prec = precision_score(trues, preds)
    rec = recall_score(trues, preds)
    roc_auc_score_value = roc_auc_score(trues, probs[:, 1]) if len(np.unique(trues)) > 1 else float("nan")

    cm = confusion_matrix(trues, preds)
    plt.figure(figsize=(6, 6))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", xticklabels=["Класс 0", "Класс 7"],
                yticklabels=["Класс 0", "Класс 7"])
    plt.title("Confusion Matrix")
    cm_path = os.path.join(save_dir, f"conf_matrix_{timestamp}.png")
    plt.savefig(cm_path)
    plt.close()

    plt.figure(figsize=(10, 6))
    plt.hist(probs[trues == 1, 1], bins=40, alpha=0.7, label="Класс 7", color="green")
    plt.hist(probs[trues == 0, 1], bins=40, alpha=0.7, label="Класс 0", color="red")
    plt.title("Распределение вероятностей")
    plt.xlabel("Вероятность быть классом 7")
    plt.legend()
    hist_path = os.path.join(save_dir, f"prob_hist_{timestamp}.png")
    plt.savefig(hist_path)
    plt.close()

    fpr, tpr, _ = roc_curve(trues, probs[:, 1])
    roc_auc = auc(fpr, tpr)
    plt.figure(figsize=(8, 6))
    plt.plot(fpr, tpr, color="orange", lw=2, label=f"ROC curve (AUC = {roc_auc:.2f})")
    plt.plot([0, 1], [0, 1], color="navy", lw=2, linestyle="--")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("ROC Curve")
    plt.legend(loc="lower right")
    roc_path = os.path.join(save_dir, f"roc_curve_{timestamp}.png")
    plt.savefig(roc_path)
    plt.close()

    doc = Document()
    doc.add_heading("Отчёт по обучению модели CNN-BiLSTM + Attention", level=1)
    doc.add_paragraph(f"Дата и время: {timestamp}")

    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Метрика"
    hdr_cells[1].text = "Значение"

    for key, value in {
        "accuracy": acc,
        "f1_score": f1,
        "precision": prec,
        "recall": rec,
        "roc_auc": roc_auc_score_value,
    }.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = f"{value:.4f}"

    doc.add_heading("Классификационный отчёт", level=2)
    doc.add_paragraph(classification_report(trues, preds, target_names=["Класс 0", "Класс 7"], digits=4))

    doc.add_heading("Примеры предсказаний", level=2)
    tab = doc.add_table(rows=1, cols=4)
    hdr = tab.rows[0].cells
    hdr[0].text = "Текст"
    hdr[1].text = "Истинная метка"
    hdr[2].text = "Прогноз"
    hdr[3].text = "Вероятность Класс 7"

    for i in range(min(10, len(trues))):
        row_cells = tab.add_row().cells
        row_cells[0].text = texts[i][:100] + "..." if len(texts[i]) > 100 else texts[i]
        row_cells[1].text = str(trues[i])
        row_cells[2].text = str(preds[i])
        row_cells[3].text = f"{probs[i, 1]:.4f}"

    doc.add_heading("Графики", level=2)
    for path in [cm_path, hist_path, roc_path]:
        if os.path.exists(path):
            doc.add_picture(path, width=Mm(120))

    doc.save(REPORT_PATH)
    logger.info(f"✔ Отчёт сохранён в {REPORT_PATH}")

    return {
        "accuracy": acc,
        "f1_score": f1,
        "precision": prec,
        "recall": rec,
        "roc_auc": roc_auc_score_value,
    }


# --- ЗАГРУЗКА И ПОДГОТОВКА ДАННЫХ ---
def load_data():
    logger.info("Загрузка и подготовка данных...")
    df = pd.read_csv(DATA_PATH, encoding="utf-8", sep=",", on_bad_lines="skip")
    assert "Класс" in df.columns and "Отзыв" in df.columns, "Нет нужных колонок!"
    df["text"] = df["Отзыв"].astype(str).apply(preproc_text)
    df["label"] = df["Класс"].apply(lambda x: 1 if x == 7 else 0)
    return df


def balance_and_split(df, max_per_class=26030, seed=SEED):   #3200,  4030
    class_1 = df[df["label"] == 1].sample(n=min(max_per_class, df["label"].sum()), random_state=seed)
    class_0 = df[df["label"] == 0].sample(n=min(max_per_class, (df["label"] == 0).sum()), random_state=seed)
    combined = pd.concat([class_1, class_0]).sample(frac=1, random_state=seed).reset_index(drop=True)
    train, test = train_test_split(combined, test_size=0.2, stratify=combined["label"], random_state=seed)
    return train, test


# --- MAIN ---
if __name__ == "__main__":
    df = load_data()
    train_df, test_df = balance_and_split(df)
    X_train = train_df["text"].tolist()
    y_train = train_df["label"].tolist()
    X_test = test_df["text"].tolist()
    y_test = test_df["label"].tolist()

    logger.info("📝 Формируем словарь и загружаем embedding матрицу...")
    word2idx, embed_matrix = build_vocab_and_embeddings(X_train, VOCAB_SIZE, EMBEDDING_PATH, EMBEDDING_DIM)

    train_dataset = ReviewDataset(X_train, y_train, word2idx)
    test_dataset = ReviewDataset(X_test, y_test, word2idx)
    train_loader = DataLoader(train_dataset, batch_size=BATCH_SIZE, shuffle=True)
    test_loader = DataLoader(test_dataset, batch_size=BATCH_SIZE)

    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = CNN_BiLSTM_Attention(
        vocab_size=len(word2idx),
        embedding_dim=EMBEDDING_DIM,
        hidden_size=HIDDEN_SIZE,
        num_classes=NUM_CLASSES,
        kernel_sizes=CNN_KERNELS,
        num_filters=CNN_FILTERS,
        dropout_rate=DROPOUT,
        embed_matrix=embed_matrix,
    ).to(device)

    criterion = nn.CrossEntropyLoss()
    optimizer = optim.Adam(model.parameters(), lr=LEARNING_RATE)

    logger.info("🏋️ Обучение...")
    for epoch in range(EPOCHS):
        train_loss = train_model(model, train_loader, criterion, optimizer, device)
        logger.info(f"Epoch {epoch + 1}/{EPOCHS}, train_loss={train_loss:.4f}")

    logger.info("📊 Оценка модели на тесте...")
    trues, preds, probs = evaluate_model(model, test_loader, device)

    logger.info("💾 Сохраняем веса модели...")
    torch.save(model.state_dict(), MODEL_PATH)
    logger.info(f"✔ Модель сохранена: {MODEL_PATH}")

    logger.info("📁 Сохраняем результаты в CSV...")
    df_results = pd.DataFrame(
        {
            "text": X_test,
            "true_label": trues,
            "predicted_label": preds,
            "prob_class_7": probs[:, 1],
            "prob_class_0": probs[:, 0],
        }
    )
    df_results.to_csv(RESULT_CSV, index=False)
    logger.info(f"✔ Результаты записаны в {RESULT_CSV}")

    logger.info("📑 Строим графики и формируем отчёт...")
    plot_and_report(trues, preds, probs, X_test, RESULTS_DIR, TIMESTAMP)
    logger.info("✅ Скрипт завершён!")
