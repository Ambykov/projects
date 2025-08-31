# Обучаем модель cointegrated/rubert-tiny2



import os
import sys
import logging
from datetime import datetime
import random
import numpy as np
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report, roc_auc_score, confusion_matrix
from sklearn.metrics import roc_curve, auc
import matplotlib.pyplot as plt
import seaborn as sns
import torch
import torch.nn.functional as F
from torch.utils.data import DataLoader

# Добавляем корень проекта в путь
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
sys.path.append(PROJECT_ROOT)

# Импорты HuggingFace
from datasets import Dataset
from transformers import BertTokenizer, BertForSequenceClassification, TrainingArguments, Trainer

# Для отчётов
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Логгирование
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("train_class_7_and_0")
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(logging.Formatter("[%(asctime)s] %(levelname)s: %(message)s"))
logger.addHandler(console_handler)

# Пути
LOGS_DIR = os.path.join(PROJECT_ROOT, "logs")
MODELS_DIR = os.path.join(PROJECT_ROOT, "models")
RESULTS_DIR = os.path.join(PROJECT_ROOT, "results")
os.makedirs(LOGS_DIR, exist_ok=True)
os.makedirs(MODELS_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)

TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M")
MODEL_SAVE_DIR = os.path.join(MODELS_DIR, f"model_class_7_and_0_{TIMESTAMP}")
RESULT_CSV = os.path.join(RESULTS_DIR, f"predictions_class_7_and_0_{TIMESTAMP}.csv")
REPORT_DOCX = os.path.join(RESULTS_DIR, f"report_class_7_and_0_{TIMESTAMP}.docx")

# Гиперпараметры
MODEL_NAME = "cointegrated/rubert-tiny2"
MAX_LENGTH = 256
BATCH_SIZE = 16
EPOCHS = 10
LEARNING_RATE = 3e-5
SEED = 42

# Фиксируем seed
random.seed(SEED)
np.random.seed(SEED)
torch.manual_seed(SEED)

def load_data():
    logger.info("📥 Загрузка данных...")
    DATA_PATH = os.path.join(PROJECT_ROOT, "data", "full_reviews_mixed.csv")
    try:
        df = pd.read_csv(DATA_PATH, sep=',', on_bad_lines='skip', encoding='utf-8', engine='python')
    except Exception as e:
        logger.error(f"❌ Ошибка при загрузке CSV: {e}")
        raise

    if 'Класс' in df.columns and 'Отзыв' in df.columns:
        df = df.rename(columns={"Класс": "label", "Отзыв": "text"})
    else:
        logger.error("❌ Колонка 'Класс' или 'Отзыв' отсутствует.")
        print(df.head(3).to_string())
        raise KeyError("Формат файла некорректен или данные повреждены.")

    df['text'] = df['text'].astype(str)
    return df

def prepare_dataset(df):
    logger.info("🛠 Подготовка датасета...")

    # Выбираем отзывы класса 7 и 0
    class_7 = df[df['label'] == 7].sample(n=min(3200, len(df[df['label'] == 7])), random_state=SEED)
    class_0 = df[df['label'] == 0].sample(n=min(3200, len(df[df['label'] == 0])), random_state=SEED)

    # Объединяем и перемешиваем
    combined = pd.concat([class_7, class_0]).sample(frac=1, random_state=SEED).reset_index(drop=True)

    # Преобразуем метки: 7 → 1, другие → 0
    combined['label'] = combined['label'].apply(lambda x: 1 if x == 7 else 0)

    logger.info(f"📊 Выборка: {len(combined)} отзывов, распределение классов:")
    print(combined['label'].value_counts())

    # Разделяем на train/test
    train_df, test_df = train_test_split(
        combined,
        test_size=0.2,
        stratify=combined['label'],
        random_state=SEED
    )

    train_dataset = Dataset.from_pandas(train_df[['text', 'label']])
    test_dataset = Dataset.from_pandas(test_df[['text', 'label']])

    return train_dataset, test_dataset

def tokenize_dataset(dataset, tokenizer):
    def tokenize_function(examples):
        texts = [str(text) for text in examples["text"]]
        return tokenizer(
            texts,
            padding="max_length",
            truncation=True,
            max_length=MAX_LENGTH,
            return_tensors="pt"
        )

    return dataset.map(tokenize_function, batched=True)

def compute_metrics(p):
    preds = p.predictions.argmax(-1)
    probs = F.softmax(torch.tensor(p.predictions), dim=-1).numpy()
    labels = p.label_ids
    return {
        "accuracy": accuracy_score(labels, preds),
        "f1": f1_score(labels, preds, average="binary"),
        "precision": precision_score(labels, preds, average="binary"),
        "recall": recall_score(labels, preds, average="binary"),
        "roc_auc": roc_auc_score(labels, probs[:, 1]),
    }

def train_model(train_dataset, eval_dataset, tokenizer):
    logger.info("🏋️ Начало обучения модели...")

    model = BertForSequenceClassification.from_pretrained(MODEL_NAME, num_labels=2)

    training_args = TrainingArguments(
        output_dir="./results",
        eval_strategy="epoch",
        logging_strategy="epoch",
        learning_rate=LEARNING_RATE,
        per_device_train_batch_size=BATCH_SIZE,
        num_train_epochs=EPOCHS,
        weight_decay=0.01,
        save_strategy="epoch",
        load_best_model_at_end=True,
        metric_for_best_model="roc_auc",
        report_to="none",
        logging_dir=LOGS_DIR,
        save_total_limit=2,
        seed=SEED,
    )

    trainer = Trainer(
        model=model,
        args=training_args,
        train_dataset=train_dataset,
        eval_dataset=eval_dataset,
        compute_metrics=compute_metrics
    )

    trainer.train()

    logger.info(f"💾 Сохраняем модель в {MODEL_SAVE_DIR}")
    model.save_pretrained(MODEL_SAVE_DIR)
    tokenizer.save_pretrained(MODEL_SAVE_DIR)

    return trainer

def evaluate_model(trainer, test_dataset):
    logger.info("📊 Оценка модели на тестовой выборке...")

    predictions = trainer.predict(test_dataset)
    probs = F.softmax(torch.tensor(predictions.predictions), dim=-1).numpy()
    preds = probs.argmax(-1)
    labels = predictions.label_ids

    # Сохраняем результаты
    df_results = pd.DataFrame({
        "text": test_dataset["text"],  # Теперь доступен, так как не удалён
        "true_label": labels,
        "predicted_label": preds,
        "prob_class_7": probs[:, 1],
        "prob_class_0": probs[:, 0]
    })

    df_results.to_csv(RESULT_CSV, index=False)
    logger.info(f"📁 Результаты сохранены в {RESULT_CSV}")

    # Метрики
    acc = accuracy_score(labels, preds)
    f1 = f1_score(labels, preds)
    prec = precision_score(labels, preds)
    rec = recall_score(labels, preds)
    roc_auc = roc_auc_score(labels, probs[:, 1])

    print(f"🎯 Метрики:")
    print(f"Accuracy: {acc:.4f}, F1: {f1:.4f}, Precision: {prec:.4f}, Recall: {rec:.4f}, ROC AUC: {roc_auc:.4f}")
    print(classification_report(labels, preds, target_names=["Класс 0", "Класс 7"]))

    return df_results, probs

def plot_results(df_results, probs, save_dir, timestamp):
    logger.info("📉 Визуализация результатов...")

    cm = confusion_matrix(df_results["true_label"], df_results["predicted_label"])
    plt.figure(figsize=(6, 6))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", xticklabels=["Класс 0", "Класс 7"],
                yticklabels=["Класс 0", "Класс 7"])
    plt.title("Confusion Matrix")
    plt.tight_layout()
    cm_path = os.path.join(save_dir, f"confusion_matrix_class_7_and_0_{timestamp}.png")
    plt.savefig(cm_path)
    plt.close()

    plt.figure(figsize=(10, 6))
    plt.hist(probs[df_results["true_label"] == 1, 1], bins=50, alpha=0.7, label="Класс 7", color="green")
    plt.hist(probs[df_results["true_label"] == 0, 1], bins=50, alpha=0.7, label="Класс 0", color="red")
    plt.title("Распределение вероятностей для классов")
    plt.xlabel("Вероятность быть классом 7")
    plt.ylabel("Частота")
    plt.legend()
    hist_path = os.path.join(save_dir, f"probability_histogram_class_7_and_0_{timestamp}.png")
    plt.savefig(hist_path)
    plt.close()

    fpr, tpr, _ = roc_curve(df_results["true_label"], probs[:, 1])
    roc_auc = auc(fpr, tpr)
    plt.figure(figsize=(8, 6))
    plt.plot(fpr, tpr, color='darkorange', lw=2, label=f'ROC curve (AUC = {roc_auc:.2f})')
    plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
    plt.xlim([0.0, 1.0])
    plt.ylim([0.0, 1.05])
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curve')
    plt.legend(loc="lower right")
    roc_path = os.path.join(save_dir, f"roc_curve_class_7_and_0_{timestamp}.png")
    plt.savefig(roc_path)
    plt.close()

    return cm_path, hist_path, roc_path

def save_report_to_docx(df_results, metrics, cm_path, hist_path, roc_path, save_dir, timestamp):
    logger.info("📄 Создание отчёта в Word (.docx)...")

    doc = Document()
    doc.add_heading("Отчёт по обучению модели: Класс 7 vs Класс 0", level=1)
    doc.add_paragraph(f"Дата и время: {timestamp}")

    doc.add_heading("Метрики качества", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Метрика"
    hdr_cells[1].text = "Значение"

    for key, value in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = f"{value:.4f}"

    doc.add_heading("Классификационный отчёт", level=2)
    report = classification_report(df_results['true_label'], df_results['predicted_label'],
                                  target_names=["Класс 0", "Класс 7"], digits=4)
    doc.add_paragraph(report)

    doc.add_heading("Примеры предсказаний", level=2)
    sample = df_results.head(10)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Текст"
    hdr_cells[1].text = "Метка"
    hdr_cells[2].text = "Предсказание"
    hdr_cells[3].text = "Вероятность (Класс 7)"
    hdr_cells[4].text = "Вероятность (Класс 0)"

    for _, row in sample.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['text'][:100] + "..." if len(row['text']) > 100 else row['text']
        row_cells[1].text = str(row['true_label'])
        row_cells[2].text = str(row['predicted_label'])
        row_cells[3].text = f"{row['prob_class_7']:.4f}"
        row_cells[4].text = f"{row['prob_class_0']:.4f}"

    doc.add_heading("Графики", level=2)
    doc.add_picture(cm_path, width=Mm(150))
    doc.add_paragraph("Матрица ошибок", style='Caption')
    doc.add_picture(hist_path, width=Mm(150))
    doc.add_paragraph("Гистограмма вероятностей", style='Caption')
    doc.add_picture(roc_path, width=Mm(150))
    doc.add_paragraph("ROC-кривая", style='Caption')

    report_path = os.path.join(save_dir, f"report_class_7_and_0_{timestamp}.docx")
    doc.save(report_path)
    logger.info(f"✅ Отчёт сохранён в {report_path}")
    return report_path

if __name__ == "__main__":
    logger.info("🚀 Начало обучения модели для классов 7 и 0")

    df = load_data()
    train_dataset, test_dataset = prepare_dataset(df)

    tokenizer = BertTokenizer.from_pretrained(MODEL_NAME)
    tokenized_train = tokenize_dataset(train_dataset, tokenizer)
    tokenized_test = tokenize_dataset(test_dataset, tokenizer).with_format("torch")

    trainer = train_model(tokenized_train, tokenized_test, tokenizer)
    df_results, probs = evaluate_model(trainer, tokenized_test)

    cm_path, hist_path, roc_path = plot_results(df_results, probs, RESULTS_DIR, TIMESTAMP)

    metrics_dict = {
        "accuracy": accuracy_score(df_results['true_label'], df_results['predicted_label']),
        "f1": f1_score(df_results['true_label'], df_results['predicted_label']),
        "precision": precision_score(df_results['true_label'], df_results['predicted_label']),
        "recall": recall_score(df_results['true_label'], df_results['predicted_label']),
        "roc_auc": roc_auc_score(df_results['true_label'], df_results['prob_class_7']),
        "loss": trainer.evaluate().get("eval_loss", 0)
    }

    save_report_to_docx(df_results, metrics_dict, cm_path, hist_path, roc_path, RESULTS_DIR, TIMESTAMP)

    logger.info("✅ Обучение и отчёт завершены успешно!")