# в этом скрипте тестируются 4600 отзывов датасета Виталия (3200 + 3200)
# в этом скрипте в файл пишется исходный текст отзыва, а не лемматизированный текст отзыва


import os
import re
import logging
from datetime import datetime
import numpy as np
import pandas as pd
import torch
from torch.utils.data import Dataset, DataLoader
import torch.nn as nn
from docx import Document
from docx.shared import Mm
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.metrics import (
    accuracy_score, f1_score, precision_score, recall_score,
    roc_auc_score, confusion_matrix, classification_report, roc_curve, auc)
import pickle

# Параметры и пути
PROJECT_ROOT = r"C:\Users\Андрей\Documents\Курсы\AI_ML_разработчик\Стажировка_2\Ноутбуки\Пайплайн"
DATA_PATH = os.path.join(PROJECT_ROOT, "data", "ALL_Reviews_P23.csv")
EMBEDDING_PATH = r"C:\Users\Андрей\Documents\Курсы\AI_ML_разработчик\Стажировка_2\ruwikiruscorpora\model.txt"

RESULTS_DIR = os.path.join(PROJECT_ROOT, "results", "cnn_bilstm_attention_test_allreviews")
os.makedirs(RESULTS_DIR, exist_ok=True)
TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M")
REPORT_PATH = os.path.join(RESULTS_DIR, f"report_test_{TIMESTAMP}.docx")
RESULT_CSV = os.path.join(RESULTS_DIR, f"predictions_test_{TIMESTAMP}.csv")
MISSED_WORDS_PATH = os.path.join(RESULTS_DIR, f"missed_words_test_{TIMESTAMP}.txt")

#WORD2IDX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_21072025_1941", "word2idx.pkl")
#EMBED_MATRIX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_21072025_1941", "embed_matrix.npy")
#MODEL_WEIGHTS_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_21072025_1941", "model_20250721_1941.pt")

#WORD2IDX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_26072025_1857", "word2idx.pkl")
#EMBED_MATRIX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_26072025_1857", "embed_matrix.npy")
#MODEL_WEIGHTS_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_26072025_1857", "model_20250726_1857.pt")

# длина 120, словарь 15 000
#WORD2IDX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1031", "word2idx.pkl")
#EMBED_MATRIX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1031", "embed_matrix.npy")
#MODEL_WEIGHTS_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1031", "model_20250727_1031.pt")

# длина 258, словарь 15 000
#WORD2IDX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1059", "word2idx.pkl")
#EMBED_MATRIX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1059", "embed_matrix.npy")
#MODEL_WEIGHTS_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1059", "model_20250727_1059.pt")

# длина 258, словарь 25 000
WORD2IDX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1216", "word2idx.pkl")
EMBED_MATRIX_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1216", "embed_matrix.npy")
MODEL_WEIGHTS_PATH = os.path.join(PROJECT_ROOT, "models", "cnn_bilstm_attention_27072025_1216", "model_20250727_1216.pt")

SEED = 42
MAX_LENGTH = 120
EMBEDDING_DIM = 300
HIDDEN_SIZE = 128
NUM_CLASSES = 2
CNN_KERNELS = (3, 4, 5)
CNN_FILTERS = 64
BATCH_SIZE = 32
DROPOUT = 0.4

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("CNN_BiLSTM_Attention_Test_AllReviews")

# Лемматизация и препроцессинг (как у вас)
import pymorphy3
morph = pymorphy3.MorphAnalyzer()
pos_map = {
    'NOUN': 'NOUN', 'VERB': 'VERB', 'INFN': 'VERB',
    'ADJF': 'ADJ', 'ADJS': 'ADJ', 'COMP': 'ADJ',
    'PRTF': 'ADJ', 'PRTS': 'ADJ', 'GRND': 'VERB',
    'NUMR': 'NUM', 'ADVB': 'ADV', 'NPRO': 'PRON',
    'PRED': 'ADV', 'PREP': 'ADP', 'CONJ': 'CONJ',
    'PRCL': 'PART', 'INTJ': 'INTJ'
}


def get_pos_tag(word):
    parse = morph.parse(word)[0]
    pos = parse.tag.POS
    return pos_map.get(pos, 'X')


def process_token(token):
    parse = morph.parse(token)[0]
    lemma = parse.normal_form
    pos = get_pos_tag(token)
    return f"{lemma}_{pos}"


def preproc_text(text):
    tokens = re.findall(r"\w+", text.lower())
    return " ".join(process_token(token) for token in tokens)


def remove_pos_tags(text):
    tokens = text.split()
    lemmas = [t.split('_')[0] if '_' in t else t for t in tokens]
    return " ".join(lemmas)


# Загрузка и подготовка тестовых данных - с отборами
def load_and_preprocess_data(sample_size_per_class=3200):
    logger.info("Загрузка исходных данных из ALL_Reviews_P23.csv...")
    df = pd.read_csv(DATA_PATH, encoding='utf-8', sep=',', on_bad_lines="skip")

    # Проверяем наличие нужных столбцов
    required_cols = ['Класс', 'Текст отзыва']
    for col in required_cols:
        if col not in df.columns:
            raise KeyError(f"В файле отсутствует колонка '{col}'")

    # Переименуем для совместимости с остальным кодом
    df = df.rename(columns={'Текст отзыва': 'Отзыв'})

    # Приводим 'Класс' к строке, если нужно, или целому числу
    df['Класс'] = df['Класс'].astype(str)

    # Фильтрация только для двух классов: '7' и остальные
    df['label'] = df['Класс'].apply(lambda x: 1 if x == '7' else 0)

    # Берём 3200 случайных отзывов класса 7
    df_pos = df[df['label'] == 1]
    df_pos_sampled = df_pos.sample(n=min(sample_size_per_class, len(df_pos)), random_state=SEED)

    # Берём 3200 случайных отзывов из других классов
    df_neg = df[df['label'] == 0]
    df_neg_sampled = df_neg.sample(n=min(sample_size_per_class, len(df_neg)), random_state=SEED)

    # Объединяем и перемешиваем
    df_sampled = pd.concat([df_pos_sampled, df_neg_sampled]).sample(frac=1, random_state=SEED).reset_index(drop=True)

    # Сохраняем исходный оригинальный текст
    df_sampled['original_text'] = df_sampled['Отзыв'].astype(str)
    # Препроцессинг
    df_sampled['text'] = df_sampled['Отзыв'].astype(str).apply(preproc_text)

    return df_sampled


# Токенизация в индекс
def text_to_sequence(text, word2idx, max_length=MAX_LENGTH):
    tokens = text.split()
    indices = [word2idx.get(t, 1) for t in tokens]  # 1 = <UNK>
    if len(indices) < max_length:
        indices += [0] * (max_length - len(indices))
    else:
        indices = indices[:max_length]
    return indices


class ReviewDataset(Dataset):
    def __init__(self, texts, labels, word2idx, max_length=MAX_LENGTH):
        self.sequences = [text_to_sequence(t, word2idx, max_length) for t in texts]
        self.labels = labels

    def __len__(self):
        return len(self.sequences)

    def __getitem__(self, idx):
        return torch.tensor(self.sequences[idx], dtype=torch.long), torch.tensor(self.labels[idx], dtype=torch.long)


# Определение модели (копия класса из вашего кода)
class CNN_BiLSTM_Attention(nn.Module):
    def __init__(self, vocab_size, embedding_dim, hidden_size, num_classes,
                 kernel_sizes=(3, 4, 5), num_filters=64, dropout_rate=0.4, embed_matrix=None):
        super().__init__()
        if embed_matrix is not None:
            self.embedding = nn.Embedding.from_pretrained(torch.FloatTensor(embed_matrix), freeze=False, padding_idx=0)
        else:
            self.embedding = nn.Embedding(vocab_size, embedding_dim, padding_idx=0)
        self.convs = nn.ModuleList([
            nn.Conv1d(embedding_dim, num_filters, k) for k in kernel_sizes
        ])
        self.lstm = nn.LSTM(num_filters * len(kernel_sizes), hidden_size, batch_first=True, bidirectional=True)
        self.dropout = nn.Dropout(dropout_rate)
        self.attn_fc = nn.Linear(hidden_size * 2, 1)
        self.fc = nn.Sequential(
            nn.Linear(hidden_size * 2, hidden_size),
            nn.ReLU(),
            nn.Dropout(dropout_rate),
            nn.Linear(hidden_size, num_classes)
        )

    def forward(self, x):
        emb = self.embedding(x)  # [batch, seq_len, emb_dim]
        emb = emb.transpose(1, 2)  # [batch, emb_dim, seq_len]
        convs = [torch.relu(conv(emb)) for conv in self.convs]
        min_len = min(conv.shape[2] for conv in convs)
        convs = [conv[:, :, :min_len] for conv in convs]
        cnn_out = torch.cat(convs, dim=1)  # [batch, num_filters*len_kernels, seq_len']
        cnn_out = cnn_out.transpose(1, 2)  # [batch, seq_len', num_filters*len_kernels]
        lstm_out, _ = self.lstm(cnn_out)  # [batch, seq_len', hidden_size*2]
        attn_weights = torch.softmax(self.attn_fc(lstm_out), dim=1)  # [batch, seq_len', 1]
        attended = torch.sum(attn_weights * lstm_out, dim=1)          # [batch, hidden_size*2]
        out = self.dropout(attended)
        return self.fc(out)


# Оценка
def evaluate_model(model, dataloader, device):
    model.eval()
    true_labels, pred_labels, probs = [], [], []
    with torch.no_grad():
        for x_batch, y_batch in dataloader:
            x_batch = x_batch.to(device)
            y_batch = y_batch.to(device)
            logits = model(x_batch)
            prob = torch.softmax(logits, dim=-1)
            preds = prob.argmax(dim=-1)
            true_labels.extend(y_batch.cpu().numpy())
            pred_labels.extend(preds.cpu().numpy())
            probs.extend(prob.cpu().numpy())
    return np.array(true_labels), np.array(pred_labels), np.array(probs)


# Визуализация и отчёт (копия вашей функции, можно импортировать, если она есть)
def plot_and_report(trues, preds, probs, original_texts, preprocessed_texts, save_dir, timestamp):
    acc = accuracy_score(trues, preds)
    f1 = f1_score(trues, preds)
    prec = precision_score(trues, preds)
    rec = recall_score(trues, preds)
    roc_auc = roc_auc_score(trues, probs[:, 1]) if len(np.unique(trues)) > 1 else float('nan')

    cm = confusion_matrix(trues, preds)
    plt.figure(figsize=(6, 6))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=['Класс 0', 'Класс 7'], yticklabels=['Класс 0', 'Класс 7'])
    plt.title('Confusion Matrix')
    cm_path = os.path.join(save_dir, f'conf_matrix_{timestamp}.png')
    plt.savefig(cm_path)
    plt.close()

    plt.figure(figsize=(10, 6))
    plt.hist(probs[trues == 1, 1], bins=40, alpha=0.7, label='Класс 7', color='green')
    plt.hist(probs[trues == 0, 1], bins=40, alpha=0.7, label='Класс 0', color='red')
    plt.xlabel('Вероятность быть классом 7')
    plt.title('Распределение вероятностей')
    plt.legend()
    hist_path = os.path.join(save_dir, f'prob_hist_{timestamp}.png')
    plt.savefig(hist_path)
    plt.close()

    fpr, tpr, _ = roc_curve(trues, probs[:, 1])
    roc_auc_val = auc(fpr, tpr)
    plt.figure(figsize=(8, 6))
    plt.plot(fpr, tpr, color='orange', lw=2, label=f'ROC curve (AUC = {roc_auc_val:.2f})')
    plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curve')
    plt.legend(loc='lower right')
    roc_path = os.path.join(save_dir, f'roc_curve_{timestamp}.png')
    plt.savefig(roc_path)
    plt.close()

    doc = Document()
    doc.add_heading('Отчёт по тестированию модели CNN-BiLSTM + Attention', level=1)
    doc.add_paragraph(f'Дата и время: {timestamp}')

    # Метрики
    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Метрика"
    hdr_cells[1].text = "Значение"
    for k, v in {
        'accuracy': acc,
        'f1_score': f1,
        'precision': prec,
        'recall': rec,
        'roc_auc': roc_auc
    }.items():
        row_cells = table.add_row().cells
        row_cells[0].text = k
        row_cells[1].text = f"{v:.4f}"

    # Классификационный отчёт
    doc.add_heading('Классификационный отчёт', level=2)
    report_text = classification_report(trues, preds, target_names=['Класс 0', 'Класс 7'], digits=4)
    doc.add_paragraph(report_text)

    # Примеры предсказаний
    doc.add_heading('Примеры предсказаний', level=2)
    sample_table = doc.add_table(rows=1, cols=4)
    hdr = sample_table.rows[0].cells
    hdr[0].text = "Текст"
    hdr[1].text = "Истинная метка"
    hdr[2].text = "Прогноз"
    hdr[3].text = "Вероятность Класс 7"

    for i in range(min(10, len(trues))):
        row_cells = sample_table.add_row().cells
        row_cells[0].text = original_texts[i][:100] + "..." if len(original_texts[i]) > 100 else original_texts[i]
        row_cells[1].text = str(trues[i])
        row_cells[2].text = str(preds[i])
        row_cells[3].text = f"{probs[i, 1]:.4f}"

    # Вставка графиков
    doc.add_heading("Графики", level=2)
    for path in [cm_path, hist_path, roc_path]:
        if os.path.exists(path):
            doc.add_picture(path, width=Mm(120))

    doc.save(REPORT_PATH)
    logger.info(f"✔ Отчёт сохранён: {REPORT_PATH}")


if __name__ == "__main__":
    logger.info("🚀 Начинаем тестирование модели CNN-BiLSTM + Attention на новых данных ALL_Reviews_P23.csv")

    # Загрузка и препроцессинг с выборками
    df_test = load_and_preprocess_data(sample_size_per_class=2300)
    original_texts_test = df_test['original_text'].tolist()
    texts_test = df_test['text'].tolist()
    labels_test = df_test['label'].tolist()

    # Загрузка словаря и эмбеддингов
    with open(WORD2IDX_PATH, 'rb') as f:
        word2idx = pickle.load(f)
    embed_matrix = np.load(EMBED_MATRIX_PATH)

    # Создаем Dataset и DataLoader
    test_dataset = ReviewDataset(texts_test, labels_test, word2idx)
    test_loader = DataLoader(test_dataset, batch_size=BATCH_SIZE)

    # Инициализация модели и загрузка весов
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model = CNN_BiLSTM_Attention(
        vocab_size=len(word2idx),
        embedding_dim=EMBEDDING_DIM,
        hidden_size=HIDDEN_SIZE,
        num_classes=NUM_CLASSES,
        kernel_sizes=CNN_KERNELS,
        num_filters=CNN_FILTERS,
        dropout_rate=DROPOUT,
        embed_matrix=embed_matrix
    )
    model.load_state_dict(torch.load(MODEL_WEIGHTS_PATH, map_location=device))
    model.to(device)

    # Оценка
    y_true, y_pred, y_probs = evaluate_model(model, test_loader, device)

    # Сохраняем результаты
    df_results = pd.DataFrame({
        "text": original_texts_test,
        "true_label": y_true,
        "predicted_label": y_pred,
        "prob_class_7": y_probs[:, 1],
        "prob_class_0": y_probs[:, 0],
    })
    df_results.to_csv(RESULT_CSV, index=False)
    logger.info(f"✔ Результаты сохранены в {RESULT_CSV}")

    # Построение графиков и отчёта
    plot_and_report(y_true, y_pred, y_probs, original_texts_test, texts_test, RESULTS_DIR, TIMESTAMP)
    logger.info("✅ Тестирование завершено")
