# src/classifiers/classifier_result.py


import os
import re
import numpy as np
import pandas as pd
from decorators import pipeline_step
from config import logger, SAVE_STEP_12_RESULT, PROCESSED_DIR, TIMESTAMP

@pipeline_step(step_number=12, step_name="КЛАССИФИКАТОР RESULT — Итоговый класс")
def classifier_result(df, stop_pipeline_flag=False, step_number=12):
    """
    Итоговый классификатор с учётом уникальных наборов классов (списков кодов),
    присвоением итогового класса по приоритетам,
    добавлением вероятностей позитивной/нейтральной тональности для класса 202,
    сохранением детализированной статистики и DOCX-отчётом с промежуточными итогами.

    Значения вероятностей сохраняются без округления,
    а доля в итоговых таблицах — с 2 знаками после запятой.
    """

    if stop_pipeline_flag:
        logger.warning(f"🔚 [{step_number}] Шаг прерван пользователем")
        return df

    if df is None or df.empty:
        logger.warning(f"🟡 [{step_number}] Входной DataFrame пустой или None")
        return df

    if "Класс" not in df.columns:
        logger.error(f"❌ [{step_number}] Входной DataFrame должен содержать колонку 'Класс'")
        return df

    rules_dict = {
        '0': (1, 1),
        '100': (1, 1),
        '999': (999, 2),
        '700': (7, 3),
        '500': (5, 4),
        '300': (3, 5),
        '200': (2, 6),
        '201': (2, 6),
        '202': (2, 6),
        '701': (7, 7),
    }

    def parse_class_codes(c):
        import numpy as np

        if isinstance(c, (list, tuple, np.ndarray)):
            return [str(x).strip() for x in c if str(x).strip() != '']

        if pd.isna(c):
            return []

        if isinstance(c, str):
            c = c.strip()
            if c.startswith("[") and c.endswith("]"):
                try:
                    import ast
                    parsed = ast.literal_eval(c)
                    if isinstance(parsed, (list, tuple)):
                        return [str(x).strip() for x in parsed if str(x).strip() != '']
                except Exception as e:
                    logger.debug(f"Не удалось распарсить строку в список: {c}, ошибка: {e}")
            if ',' in c:
                parts = c.split(',')
                return [p.strip() for p in parts if p.strip() != '']
            return [c] if c else []

        if hasattr(c, 'values'):
            return [str(x).strip() for x in c.values.flatten() if str(x).strip() != '']

        if hasattr(c, '__iter__') and not isinstance(c, (str, bytes)):
            try:
                return [str(x).strip() for x in c if str(x).strip() != '']
            except Exception:
                pass

        if isinstance(c, (int, float)):
            return [str(int(c))] if int(c) == c else [str(c)]

        return [str(c).strip()]

    def parse_probs_from_note(note):
        if not isinstance(note, str):
            return (np.nan, np.nan)
        pos_match = re.search(r'pos=([0-9]*\.?[0-9]+)', note)
        neu_match = re.search(r'neu=([0-9]*\.?[0-9]+)', note)
        pos_prob = float(pos_match.group(1)) if pos_match else np.nan
        neu_prob = float(neu_match.group(1)) if neu_match else np.nan
        return pos_prob, neu_prob

    df = df.copy()

    if ('prob_positive_202' not in df.columns or 'prob_neutral_202' not in df.columns) or \
       df['prob_positive_202'].isna().all() or df['prob_neutral_202'].isna().all():
        probs = df['Примечание'].apply(parse_probs_from_note)
        df['prob_positive_202'] = [p[0] for p in probs]
        df['prob_neutral_202'] = [p[1] for p in probs]

    for col in ['prob_positive_202', 'prob_neutral_202']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.replace(',', '.').astype(float, errors='ignore')

    df['class_codes_list'] = df['Класс'].apply(parse_class_codes)
    df['class_codes_tuple'] = df['class_codes_list'].apply(lambda codes: tuple(sorted(set(codes))))

    def assign_result_class_by_codes(codes_tuple):
        candidates = []
        for code in codes_tuple:
            if code in rules_dict:
                candidates.append(rules_dict[code])
        if not candidates:
            return None
        candidates.sort(key=lambda x: x[1])
        return candidates[0][0]

    df['Итоговый класс'] = df['class_codes_tuple'].apply(assign_result_class_by_codes)

    total_count = len(df)

    summary_data = []
    classes_present = df['Итоговый класс'].dropna().unique()
    for cls in sorted(classes_present, key=lambda x: (float('inf') if x is None else x)):
        sub_df = df[df['Итоговый класс'] == cls]
        count = len(sub_df)
        share = count / total_count if total_count > 0 else 0.0
        codes_lists = sub_df['Класс'].apply(parse_class_codes)
        unique_codes = set(code for clist in codes_lists for code in clist)

        if cls == 202:
            avg_prob_pos = sub_df['prob_positive_202'].mean(skipna=True)
            avg_prob_neu = sub_df['prob_neutral_202'].mean(skipna=True)
        else:
            avg_prob_pos = None
            avg_prob_neu = None

        summary_data.append({
            "Итоговый класс": cls,
            "Количество отзывов": count,
            "Доля от общего": round(share, 4),  # округляем до 4 знаков для внутренней точности, но в выводе округлим до 2
            "Присутствующие коды 'Класс'": ", ".join(sorted(unique_codes)),
            "Средняя вероятность позитивной 202": avg_prob_pos if avg_prob_pos is not None else "",
            "Средняя вероятность нейтральной 202": avg_prob_neu if avg_prob_neu is not None else "",
        })
    summary_df = pd.DataFrame(summary_data)

    detail_group = (
        df.groupby(['Итоговый класс', 'class_codes_tuple'])
        .size()
        .reset_index(name='Количество отзывов')
    )
    detail_group['Доля'] = detail_group['Количество отзывов'] / total_count
    detail_group['Класс'] = detail_group['class_codes_tuple'].apply(lambda tup: "[" + ", ".join(tup) + "]")
    detail_group = detail_group[['Итоговый класс', 'Класс', 'Количество отзывов', 'Доля']]

    summary_row = pd.DataFrame({
        'Итоговый класс': ['Итого:'],
        'Класс': ['-'],
        'Количество отзывов': [detail_group['Количество отзывов'].sum()],
        'Доля': [detail_group['Доля'].sum()]
    })
    detail_group = pd.concat([detail_group, summary_row], ignore_index=True)

    for col in ['class_codes_list', 'class_codes_tuple']:
        if col in df.columns:
            df.drop(columns=[col], inplace=True)

    if SAVE_STEP_12_RESULT:
        result_file = os.path.join(PROCESSED_DIR, f"step_{step_number}_result_classification_{TIMESTAMP}.xlsx")
        docx_file = os.path.join(PROCESSED_DIR, f"step_{step_number}_result_classification_summary_{TIMESTAMP}.docx")

        try:
            with pd.ExcelWriter(result_file) as writer:
                # Сохраняем колонки с полной точностью вероятностей, долю выводим с точностью 2 знака
                df.to_excel(writer, index=False, sheet_name="Отзывы с итоговым классом")

                # Округляем Долю до 2 знаков для summary_df (только визуально)
                summary_df_to_save = summary_df.copy()
                summary_df_to_save['Доля от общего'] = summary_df_to_save['Доля от общего'].round(2)
                summary_df_to_save.to_excel(writer, index=False, sheet_name="Статистика итоговых классов")

                # Для детализации также округляем Доля при сохранении
                detail_group_to_save = detail_group.copy()
                detail_group_to_save['Доля'] = detail_group_to_save['Доля'].round(2)
                detail_group_to_save.to_excel(writer, index=False, sheet_name="Детализация по классам")

            logger.info(f"📌 [{step_number}]: Итоговые результаты сохранены в файл {result_file}")
        except Exception as e:
            logger.warning(f"⚠️ [{step_number}]: Ошибка при сохранении Excel файла: {e}")

        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            doc = Document()

            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            rFonts = style.element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), 'Calibri')

            heading = doc.add_heading('Итоговая классификация отзывов', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Дата и время выполнения: {TIMESTAMP}")
            doc.add_paragraph(f"Всего отзывов: {total_count}")
            doc.add_paragraph("")

            table = doc.add_table(rows=1, cols=4)
            table.style = 'LightShading-Accent1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Итоговый класс'
            hdr_cells[1].text = 'Класс'
            hdr_cells[2].text = 'Количество отзывов'
            hdr_cells[3].text = 'Доля'

            detail_no_total = detail_group[detail_group['Итоговый класс'] != 'Итого:']

            for group_key, group_df in detail_no_total.groupby('Итоговый класс'):
                for _, row in group_df.iterrows():
                    cells = table.add_row().cells
                    cells[0].text = str(row['Итоговый класс'])
                    cells[1].text = str(row['Класс'])
                    cells[2].text = str(row['Количество отзывов'])
                    # округляем до 2 знаков для вида в docx
                    cells[3].text = f"{row['Доля']*100:.2f}%"

                    for idx, cell in enumerate(cells):
                        for paragraph in cell.paragraphs:
                            if idx < 2:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                total_reviews = group_df['Количество отзывов'].sum()
                total_share = group_df['Доля'].sum()
                cells = table.add_row().cells
                cells[0].text = f"всего по классу {group_key}"
                cells[1].text = ""
                cells[2].text = str(total_reviews)
                cells[3].text = f"{total_share*100:.2f}%"

                for idx, cell in enumerate(cells):
                    for paragraph in cell.paragraphs:
                        paragraph.runs[0].font.bold = True
                        if idx < 2:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            total_row = detail_group[detail_group['Итоговый класс'] == 'Итого:'].iloc[0]
            cells = table.add_row().cells
            cells[0].text = str(total_row['Итоговый класс'])
            cells[1].text = str(total_row['Класс'])
            cells[2].text = str(total_row['Количество отзывов'])
            cells[3].text = f"{total_row['Доля']*100:.2f}%"

            for idx, cell in enumerate(cells):
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].font.bold = True
                    if idx < 2:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            doc.save(docx_file)
            logger.info(f"📌 [{step_number}]: Итоговая таблица классификации с промежуточными итогами сохранена в DOCX-файл {docx_file}")
        except Exception as e:
            logger.warning(f"⚠️ [{step_number}]: Ошибка при сохранении DOCX файла: {e}")

    logger.info(f"🔢 [{step_number}] Итоговая статистика по классам:")
    for _, row in summary_df.iterrows():
        pos_prob_str = (
            f", средняя вероятность позитивной 202: {row['Средняя вероятность позитивной 202']}"
            if row.get('Средняя вероятность позитивной 202') != "" else ""
        )
        neu_prob_str = (
            f", средняя вероятность нейтральной 202: {row['Средняя вероятность нейтральной 202']}"
            if row.get('Средняя вероятность нейтральной 202') != "" else ""
        )
        logger.info(
            f"Класс {row['Итоговый класс']}: {row['Количество отзывов']} отзывов, "
            f"доля {row['Доля от общего']:.2f}, коды: {row['Присутствующие коды \'Класс\'']}"
            + pos_prob_str + neu_prob_str
        )

    logger.info(f"🔢 [{step_number}] Детализация по итоговым классам и исходным кодам:")
    for _, row in detail_group.iterrows():
        cls = row['Итоговый класс']
        code = row['Класс']
        count = row['Количество отзывов']
        share_pct = row['Доля'] * 100
        if str(cls).lower() == 'итого:':
            continue
        logger.info(f"Итоговый класс {cls}, Класс {code}: {count} отзывов, доля {share_pct:.2f}%")

    return df
